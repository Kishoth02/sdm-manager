"""
file_manager.py
---------------
Manages the MCQ library folder.
"""

import json
import os
import uuid
from datetime import datetime
from typing import Optional

import httpx

GENERIC_FILENAMES = {"mcq", "questions", "test", "quiz", "file", "data", "sheet", "upload", "sample", "document", "untitled"}

def _is_generic_filename(name: str) -> bool:
    base = os.path.splitext(os.path.basename(name))[0].lower().strip()
    first_word = base.split("_")[0]
    return base in GENERIC_FILENAMES or first_word in GENERIC_FILENAMES

def _topic_from_filename(name: str) -> str:
    base = os.path.splitext(os.path.basename(name))[0]
    return base.split("_")[0].title()

def _detect_topic_llama(question_text: str, options: str = "") -> str:
    try:
        prompt = (
            f"You are an expert MCQ topic classifier.\n"
            f"Look at the question AND the options carefully.\n"
            f"Reply with ONLY one short topic name from this list: "
            f"Python, Java, C, C++, JavaScript, Web Tech, Database, SQL, Networking, OS, Data Structures, Algorithms, Aptitude, Math, English.\n"
            f"Rules:\n"
            f"- If question involves Python syntax/functions/keywords → Python\n"
            f"- If question involves Java syntax/keywords/OOP → Java\n"
            f"- If options contain code snippets → identify the language\n"
            f"- Never reply with 'General' or 'Programming'\n"
            f"- Reply with ONE word only. No punctuation. No explanation.\n\n"
            f"Question: {question_text}\nOptions: {options}\n\nTopic:"
        )
        resp = httpx.post(
            "http://localhost:11434/api/generate",
            json={"model": "mcq-topic", "prompt": prompt, "stream": False},
            timeout=15
        )
        result = resp.json().get("response", "").strip().split("\n")[0].strip().title()
        return result.title() if result else ""
    except Exception:
        return ""

BASE_DIR    = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
LIBRARY_DIR = os.path.join(BASE_DIR, "library")
os.makedirs(LIBRARY_DIR, exist_ok=True)

_index:  dict          = {}
_active: Optional[str] = None

def _detect_topics_batch(mcqs: list) -> list:
    results = []
    for m in mcqs:
        try:
            opts = ", ".join(filter(None, [m.get("option_a",""), m.get("option_b",""), m.get("option_c",""), m.get("option_d","")]))
            prompt = f"Question: {m.get('question','')}\nOptions: {opts}"
            resp = httpx.post(
                "http://localhost:11434/api/generate",
                json={"model": "mcq-topic", "prompt": prompt, "stream": False},
                timeout=15
            )
            result = resp.json().get("response", "").strip().split("\n")[0].strip()
            VALID_TOPICS = {"Python", "Java", "C", "C++", "JavaScript", "SQL", "Database", "Networking", "OS", "Web", "Aptitude", "Math", "Data Structures"}
            if result not in VALID_TOPICS:
                result = ""
            print(f"[TOPIC DEBUG] Q: {m.get('question','')[:50]} → {result}")
            results.append(result)
        except Exception:
            results.append("")
    return results


# ── Normalise one raw MCQ row ────────────────────────────────────────────────

def _norm(raw: dict) -> dict:
    def g(*keys):
        for k in keys:
            for variant in (k, k.lower(), k.upper(), k.capitalize()):
                v = raw.get(variant)
                if v is not None and str(v).strip():
                    return str(v).strip()
        return ""

    return {
        "id":          str(uuid.uuid4())[:8],
        "question":    g("question"),
        "option_a":    g("option_a", "A", "a"),
        "option_b":    g("option_b", "B", "b"),
        "option_c":    g("option_c", "C", "c"),
        "option_d":    g("option_d", "D", "d"),
        "answer":      (g("answer") or "A").upper()[:1],
        "difficulty":  g("difficulty").capitalize() if g("difficulty") else "",
        "topic":       g("topic").title() if g("topic") else "",
        "tags":        g("tags"),
        "explanation": g("explanation"),
    }


# ── Format parsers ───────────────────────────────────────────────────────────

def _parse_json(content: bytes) -> list:
    data = json.loads(content)
    if "SDM" in data:
        return _parse_sdm(data)
    rows = data if isinstance(data, list) else data.get("mcqs", [])
    return [_norm(r) for r in rows if isinstance(r, dict)]


def _parse_sdm(data: dict) -> list:
    try:
        sdm = data["SDM"]["Building"]["TemplateVersion"][0]
    except (KeyError, IndexError):
        return []

    item_defs = {item["OID"]: item for item in sdm.get("ItemDef", [])}

    code_lists = {}
    for cl in sdm.get("CodeList", []):
        oid = cl.get("OID", "")
        items = cl.get("CodeListItem", []) or cl.get("EnumeratedItem", [])
        options = []
        for ci in items:
            text = ""
            if "Decode" in ci:
                text = ci["Decode"].get("TranslatedText", {}).get("text", "")
            elif "TranslatedText" in ci:
                text = ci["TranslatedText"].get("text", "")
            val = ci.get("CodedValue", "")
            options.append({"value": val, "text": text or val})
        code_lists[oid] = options

    condition_defs = {}
    for cond in sdm.get("ConditionDef", []):
        oid = cond.get("OID", "")
        expr = cond.get("FormalExpression", {}).get("text", "")
        condition_defs[oid] = expr

    mcqs = []
    for ig in sdm.get("ItemGroupDef", []):
        group_oid = ig.get("OID", "")
        desc = ig.get("Description", {})
        desc_text = ""
        if isinstance(desc, dict):
            tt = desc.get("TranslatedText", {})
            if isinstance(tt, dict):
                desc_text = tt.get("text", "")

        raw_name = ig.get("Name", "") or ig.get("name", "") or ""
        if desc_text:
            group_name = desc_text
        elif raw_name and raw_name != group_oid:
            group_name = raw_name
        else:
            group_name = f"Group {group_oid.replace('IG.', '')}"

        for item_ref in ig.get("ItemRef", []):
            item_oid = item_ref.get("ItemOID", "")
            item = item_defs.get(item_oid, {})
            if not item:
                continue

            q_text = ""
            q_node = item.get("Question", {})
            if isinstance(q_node, dict):
                tt = q_node.get("TranslatedText", {})
                if isinstance(tt, dict):
                    q_text = tt.get("text", "")
                elif isinstance(tt, list):
                    q_text = tt[0].get("text", "") if tt else ""

            if not q_text:
                continue

            cl_ref = item.get("CodeListRef", {})
            cl_oid = cl_ref.get("CodeListOID", "") if isinstance(cl_ref, dict) else ""
            options = code_lists.get(cl_oid, [])

            opt_map = {}
            opt_values = {}
            for i, opt in enumerate(options[:4]):
                key = ["option_a", "option_b", "option_c", "option_d"][i]
                opt_map[key] = opt.get("text", opt.get("value", ""))
                opt_values[key] = opt.get("value", "")

            cond_oid = item_ref.get("CollectionExceptionConditionOID", "")
            condition_expr = condition_defs.get(cond_oid, "") if cond_oid else ""
            condition = None
            if condition_expr:
                import re
                m = re.match(r'^(!?)\s*\(\s*(\S+)\s*==\s*["\']?([^"\']+)["\']?\s*\)$', condition_expr.strip())
                if m:
                    is_negated = m.group(1) == "!"
                    condition = {
                        "parent_item_oid": m.group(2),
                        "parent_value": m.group(3),
                        "hide_when_match": is_negated,
                    }

            mcq = {
                "id": str(__import__('uuid').uuid4())[:8],
                "question": q_text,
                "option_a": opt_map.get("option_a", ""),
                "option_b": opt_map.get("option_b", ""),
                "option_c": opt_map.get("option_c", ""),
                "option_d": opt_map.get("option_d", ""),
                "option_values": opt_values,
                "answer": opt_map.get("option_a", ""),
                "difficulty": "",
                "topic": group_name,
                "tags": item.get("DataType", ""),
                "explanation": item_oid,
                "item_oid": item_oid,
                "group_oid": group_oid,
                "condition": condition,
            }
            mcqs.append(mcq)

    return mcqs


def _parse_csv(content: bytes) -> list:
    import pandas as pd, io
    df = pd.read_csv(io.BytesIO(content))
    df.columns = [c.strip().lower().replace(" ", "_") for c in df.columns]
    return [_norm(r) for r in df.to_dict(orient="records")]


def _parse_excel(content: bytes) -> list:
    import pandas as pd, io
    df = pd.read_excel(io.BytesIO(content))
    df.columns = [c.strip().lower().replace(" ", "_") for c in df.columns]
    return [_norm(r) for r in df.to_dict(orient="records")]


def _parse_txt(content: bytes) -> list:
    import re
    text = content.decode("utf-8", errors="ignore")
    lines = text.splitlines()

    cleaned = []
    for line in lines:
        line = line.strip()
        if re.match(r'^[📘📗📙📕🔷🔹▶►•\-=]{1,3}\s*(section|part|chapter|unit)', line, re.IGNORECASE):
            continue
        if re.match(r'^[─═\-=]{4,}', line):
            continue
        cleaned.append(line)
    full_text = "\n".join(cleaned)

    block_pattern = re.compile(r'(?=^(?:Q\d+[\.\)]\s*|\d+[\.\)]\s*))', re.MULTILINE)
    blocks = block_pattern.split(full_text)
    if len(blocks) <= 1:
        block_pattern2 = re.compile(r'(?=^Question\s*:)', re.MULTILINE | re.IGNORECASE)
        blocks = block_pattern2.split(full_text)

    blocks = [b.strip() for b in blocks if b.strip()]
    mcqs = []
    for block in blocks:
        mcq = _parse_block(block)
        if mcq:
            mcqs.append(_norm(mcq))
    return mcqs


def _parse_pdf(content: bytes) -> list:
    import pdfplumber, io, re
    text = ""
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            text += (page.extract_text() or "") + "\n"

    diff_map = {}
    current_diff = None
    for line in text.splitlines():
        line = line.strip()
        m = re.match(r'^(Easy|Medium|Hard)\s+Level\s*$', line, re.IGNORECASE)
        if m:
            current_diff = m.group(1).capitalize()
            continue
        qnum = re.match(r'^(\d+)[\.\)]\s*', line)
        if qnum and current_diff:
            diff_map[int(qnum.group(1))] = current_diff

    def split_inline_options(line):
        parts = re.split(r'\s+(?=[A-Da-d][\)\.])', line)
        return '\n'.join(parts) if len(parts) > 1 else line

    lines = text.splitlines()
    new_lines = []
    for line in lines:
        line = line.strip()
        if re.match(r'^(Easy|Medium|Hard)\s+Level\s*$', line, re.IGNORECASE):
            continue
        if re.search(r'[A-Da-d][\)\.]\s+\S.*[A-Da-d][\)\.]\s*', line):
            new_lines.append(split_inline_options(line))
        else:
            new_lines.append(line)

    fixed_text = '\n'.join(new_lines)
    fixed_text = re.sub(
        r'(Answer\s*:\s*)([A-Da-d])[\)\.]\s*\S+',
        lambda m: m.group(1) + m.group(2).upper(),
        fixed_text, flags=re.IGNORECASE
    )

    mcqs = _parse_txt(fixed_text.encode("utf-8"))
    diff_values = [diff_map[k] for k in sorted(diff_map.keys())]
    for i, mcq in enumerate(mcqs):
        if i < len(diff_values):
            mcq["difficulty"] = diff_values[i]

    return mcqs


def _parse_block(block: str) -> dict:
    import re
    lines = [l.strip() for l in block.splitlines() if l.strip()]
    if not lines:
        return None

    raw = {}
    q_text = re.sub(r'^(?:Q\d+[\.\)]\s*|\d+[\.\)]\s*|Question\s*:\s*)', '', lines[0], flags=re.IGNORECASE).strip()
    if not q_text:
        return None
    raw["question"] = q_text
    code_lines = []

    opt_pat  = re.compile(r'^[\(\[]?([A-Da-d1-4])[\)\]\.]\s*(.*)', re.IGNORECASE)
    ans_pat  = re.compile(r'^(?:answer|ans|correct(?:\s+answer)?)\s*[:\-]\s*(.+)', re.IGNORECASE)
    diff_pat = re.compile(r'^difficulty\s*[:\-]\s*(.+)', re.IGNORECASE)
    top_pat  = re.compile(r'^topic\s*[:\-]\s*(.+)', re.IGNORECASE)
    expl_pat = re.compile(r'^explanation\s*[:\-]\s*(.+)', re.IGNORECASE)
    letter_map = {'1':'A','2':'B','3':'C','4':'D'}
    opt_map = {}

    for line in lines[1:]:
        m = opt_pat.match(line)
        if not m and not ans_pat.match(line) and not diff_pat.match(line) and not top_pat.match(line) and not expl_pat.match(line):
            code_lines.append(line)
            continue
        if m:
            key = letter_map.get(m.group(1).upper(), m.group(1).upper())
            opt_map[key] = m.group(2).strip()
            continue
        m = ans_pat.match(line)
        if m:
            ans_raw = m.group(1).strip()
            letter = re.match(r'^([A-Da-d1-4])[\)\.\s]', ans_raw)
            if letter:
                key = letter.group(1).upper()
                raw["answer"] = letter_map.get(key, key)
            else:
                raw["answer_text"] = ans_raw
            continue
        m = diff_pat.match(line)
        if m: raw["difficulty"] = m.group(1).strip(); continue
        m = top_pat.match(line)
        if m: raw["topic"] = m.group(1).strip(); continue
        m = expl_pat.match(line)
        if m: raw["explanation"] = m.group(1).strip(); continue

    if code_lines:
        raw["question"] = raw["question"] + " " + " ".join(code_lines)
    raw["option_a"] = opt_map.get("A", "")
    raw["option_b"] = opt_map.get("B", "")
    raw["option_c"] = opt_map.get("C", "")
    raw["option_d"] = opt_map.get("D", "")

    if "answer" not in raw and "answer_text" in raw:
        ans_text = raw.pop("answer_text").lower()
        for k, v in opt_map.items():
            if ans_text in v.lower() or v.lower() in ans_text:
                raw["answer"] = k
                break
        else:
            raw["answer"] = "A"
    elif "answer" not in raw:
        raw["answer"] = "A"

    filled = sum(1 for k in ["option_a","option_b","option_c","option_d"] if raw.get(k))
    if not raw.get("question") or filled < 2:
        return None
    return raw


# ── Disk helpers ─────────────────────────────────────────────────────────────

def _write_library(fname: str):
    path = os.path.join(LIBRARY_DIR, fname)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(_index[fname], f, indent=2, ensure_ascii=False)


def _write_original(fname: str):
    src = _index[fname]["meta"].get("original_path", "")
    if not src or not os.path.isfile(src):
        return
    mcqs = _index[fname]["mcqs"]
    ext  = os.path.splitext(src)[1].lower()
    try:
        if ext == ".json":
            try:
                with open(src, "r", encoding="utf-8") as f:
                    original = json.load(f)
                if "SDM" in original:
                    sdm = original["SDM"]["Building"]["TemplateVersion"][0]
                    item_defs = {item["OID"]: item for item in sdm.get("ItemDef", [])}
                    code_lists = {cl["OID"]: cl for cl in sdm.get("CodeList", [])}
                    for mcq in mcqs:
                        item_oid = mcq.get("explanation", "")
                        if not item_oid or item_oid not in item_defs:
                            continue
                        cl_ref = item_defs[item_oid].get("CodeListRef", {})
                        cl_oid = cl_ref.get("CodeListOID", "") if isinstance(cl_ref, dict) else ""
                        if not cl_oid or cl_oid not in code_lists:
                            continue
                        new_items = []
                        for i, key in enumerate(["option_a","option_b","option_c","option_d"], 1):
                            val = mcq.get(key, "")
                            if val and val.strip():
                                new_items.append({
                                    "CodedValue": str(i),
                                    "IsEnabled": "Yes",
                                    "Decode": {"TranslatedText": {"lang": "en", "text": val}}
                                })
                        if new_items:
                            code_lists[cl_oid]["CodeListItem"] = new_items
                    with open(src, "w", encoding="utf-8") as f:
                        json.dump(original, f, indent=2, ensure_ascii=False)
                    return
            except Exception:
                pass
            with open(src, "w", encoding="utf-8") as f:
                json.dump(mcqs, f, indent=2, ensure_ascii=False)
        elif ext == ".csv":
            import pandas as pd
            pd.DataFrame(mcqs).to_csv(src, index=False)
        elif ext in (".xlsx", ".xls"):
            import pandas as pd
            pd.DataFrame(mcqs).to_excel(src, index=False)
        elif ext == ".txt":
            lines = []
            for m in mcqs:
                lines += [
                    f"Question: {m['question']}",
                    f"A) {m['option_a']}", f"B) {m['option_b']}",
                    f"C) {m['option_c']}", f"D) {m['option_d']}",
                    f"Answer: {m['answer']}", f"Difficulty: {m['difficulty']}",
                    f"Topic: {m['topic']}", f"Tags: {m['tags']}",
                    f"Explanation: {m['explanation']}", ""
                ]
            with open(src, "w", encoding="utf-8") as f:
                f.write("\n".join(lines))
    except Exception:
        pass


def _save(fname: str):
    _index[fname]["meta"]["count"]      = len(_index[fname]["mcqs"])
    _index[fname]["meta"]["updated_at"] = datetime.now().isoformat()
    _write_library(fname)
    _write_original(fname)


# ── Startup ──────────────────────────────────────────────────────────────────

def startup():
    global _active
    for fname in sorted(os.listdir(LIBRARY_DIR)):
        if fname.endswith(".json"):
            path = os.path.join(LIBRARY_DIR, fname)
            try:
                with open(path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                if "mcqs" in data:
                    _index[fname] = data
            except Exception:
                pass
    if _index and _active is None:
        _active = next(iter(_index))


# ── Public: file operations ──────────────────────────────────────────────────

def ingest(filename: str, content: bytes, original_path: str = "") -> dict:
    global _active

    base      = os.path.splitext(os.path.basename(filename))[0]
    json_name = base + ".json"

    if json_name in _index:
        _active = json_name
        existing_count = len(_index[json_name]["mcqs"])
        return {"filename": json_name, "count": existing_count}

    lib_path = os.path.join(LIBRARY_DIR, json_name)
    if os.path.exists(lib_path):
        with open(lib_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        if "mcqs" in data:
            _index[json_name] = data
            _active = json_name
            return {"filename": json_name, "count": len(data["mcqs"])}

    ext = os.path.splitext(filename)[1].lower().lstrip(".")
    if ext == "json":
        mcqs = _parse_json(content)
    elif ext == "csv":
        mcqs = _parse_csv(content)
    elif ext in ("xlsx", "xls"):
        mcqs = _parse_excel(content)
    elif ext == "txt":
        mcqs = _parse_txt(content)
    elif ext == "pdf":
        mcqs = _parse_pdf(content)
    else:
        raise ValueError(f"Unsupported format: .{ext}  (use json / csv / xlsx / txt)")

    if not mcqs:
        raise ValueError("No valid MCQs found. Check your file structure.")

    file_topic = "" if _is_generic_filename(filename) else _topic_from_filename(filename)

    if not file_topic:
        detected = _detect_topics_batch(mcqs)
        for i, m in enumerate(mcqs):
            m["topic"] = detected[i] if i < len(detected) else ""

    for m in mcqs:
        m["id"] = str(uuid.uuid4())[:8]
        if not m.get("topic"):
            m["topic"] = file_topic

    _index[json_name] = {
        "meta": {
            "filename":          json_name,
            "original_filename": filename,
            "original_path":     original_path,
            "uploaded_at":       datetime.now().isoformat(),
            "count":             len(mcqs),
        },
        "mcqs": mcqs,
    }
    _write_library(json_name)
    _active = json_name
    return {"filename": json_name, "count": len(mcqs)}


def list_files() -> list:
    return [
        {
            "filename":    fname,
            "original":    d["meta"].get("original_filename", fname),
            "count":       d["meta"]["count"],
            "uploaded_at": d["meta"]["uploaded_at"],
            "active":      fname == _active,
        }
        for fname, d in _index.items()
    ]


def switch(filename: str) -> str:
    global _active
    if filename not in _index:
        matches = [f for f in _index if f.startswith(filename.replace(".json", ""))]
        if not matches:
            raise ValueError(f"File '{filename}' not in library.")
        filename = matches[0]
    _active = filename
    return filename


def remove(filename: str) -> str:
    global _active
    if filename not in _index:
        raise ValueError(f"File '{filename}' not found.")
    del _index[filename]
    _active = next(iter(_index)) if _index else None
    return filename


def get_active() -> Optional[str]:
    return _active


# ── Public: querying ─────────────────────────────────────────────────────────

def query(
    filename:   Optional[str] = None,
    difficulty: Optional[str] = None,
    topic:      Optional[str] = None,
    search:     Optional[str] = None,
    limit:      int           = 50,
    offset:     int           = 0,
    cross_file: bool          = False,
) -> list:
    if cross_file:
        sources = list(_index.values())
    else:
        target = filename or _active
        if not target or target not in _index:
            return []
        sources = [_index[target]]

    results = []
    for entry in sources:
        fname = entry["meta"]["filename"]
        for mcq in entry["mcqs"]:
            if difficulty and mcq.get("difficulty", "").lower() != difficulty.lower():
                continue
            if topic and topic.lower() not in mcq.get("topic", "").lower():
                continue
            if search and search.lower() not in mcq.get("question", "").lower():
                continue
            results.append({**mcq, "source_file": fname})

    return results[offset: offset + limit]


def stats() -> list:
    out = []
    for fname, entry in _index.items():
        mcqs   = entry["mcqs"]
        diffs  = {}
        topics = set()
        for m in mcqs:
            d = m.get("difficulty", "Unknown")
            diffs[d] = diffs.get(d, 0) + 1
            t = m.get("topic", "")
            if t:
                topics.add(t)
        out.append({
            "filename":     fname,
            "total":        len(mcqs),
            "difficulties": diffs,
            "topics":       sorted(topics),
            "active":       fname == _active,
        })
    return out


# ── Public: editing ──────────────────────────────────────────────────────────

def add_mcq(data: dict, filename: Optional[str] = None, position: Optional[int] = None) -> dict:
    target = filename or _active
    if not target or target not in _index:
        raise ValueError("No active file. Load a file first.")

    new_q = data.get("question", "").strip().lower()
    for existing in _index[target]["mcqs"]:
        if existing.get("question", "").strip().lower() == new_q:
            raise ValueError(f"This question already exists in the file!")

    if data.get("topic"):
        data["topic"] = data["topic"].strip().title()
    data["id"] = str(uuid.uuid4())[:8]
    mcq = _norm(data)
    if not mcq.get("topic"):
        opts = " ".join([mcq.get("option_a",""), mcq.get("option_b",""), mcq.get("option_c",""), mcq.get("option_d","")])
        mcq["topic"] = _detect_topic_llama(mcq.get("question",""), opts)

    topic = (mcq.get("topic") or "").strip().lower()

    if position is not None:
        if topic:
            topic_qs = [i for i, m in enumerate(_index[target]["mcqs"])
                        if (m.get("topic") or "").strip().lower() == topic]
            if topic_qs and position <= len(topic_qs):
                actual_idx = topic_qs[position - 2] + 1
                _index[target]["mcqs"].insert(actual_idx, mcq)
            else:
                idx = max(0, min(position - 1, len(_index[target]["mcqs"])))
                _index[target]["mcqs"].insert(idx, mcq)
        else:
            idx = max(0, min(position - 1, len(_index[target]["mcqs"])))
            _index[target]["mcqs"].insert(idx, mcq)
    else:
        if topic:
            last_topic_idx = None
            for i, m in enumerate(_index[target]["mcqs"]):
                if (m.get("topic") or "").strip().lower() == topic:
                    last_topic_idx = i
            if last_topic_idx is not None:
                _index[target]["mcqs"].insert(last_topic_idx + 1, mcq)
            else:
                _index[target]["mcqs"].append(mcq)
        else:
            _index[target]["mcqs"].append(mcq)

    _save(target)
    return mcq


def delete_mcq(mcq_id: str, filename: Optional[str] = None) -> bool:
    target = filename or _active
    if not target or target not in _index:
        raise ValueError("No active file.")
    before = len(_index[target]["mcqs"])
    _index[target]["mcqs"] = [m for m in _index[target]["mcqs"] if m["id"] != mcq_id]
    changed = len(_index[target]["mcqs"]) < before
    if changed:
        _save(target)
    return changed


def edit_mcq(mcq_id: str, updates: dict, filename: Optional[str] = None) -> Optional[dict]:
    target = filename or _active
    if not target or target not in _index:
        raise ValueError("No active file.")
    for mcq in _index[target]["mcqs"]:
        if mcq["id"] == mcq_id:
            mcq.update({k: v for k, v in updates.items() if k != "id"})
            _save(target)
            return mcq
    return None


def export(filename: Optional[str] = None, fmt: str = "json") -> dict:
    target = filename or _active
    if not target or target not in _index:
        raise ValueError("No active file.")
    mcqs = _index[target]["mcqs"]

    if fmt == "json":
        return {"content": json.dumps(mcqs, indent=2, ensure_ascii=False),
                "mime": "application/json", "ext": "json"}

    if fmt == "csv":
        import pandas as pd, io
        buf = io.StringIO()
        pd.DataFrame(mcqs).to_csv(buf, index=False)
        return {"content": buf.getvalue(), "mime": "text/csv", "ext": "csv"}

    if fmt == "xlsx":
        import pandas as pd, io
        buf = io.BytesIO()
        df = pd.DataFrame(mcqs)
        with pd.ExcelWriter(buf, engine="openpyxl") as writer:
            df.to_excel(writer, index=False, sheet_name="MCQs")
        return {"content": buf.getvalue(), "mime": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "ext": "xlsx"}

    if fmt == "pdf":
        import io
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4,
                                leftMargin=2*cm, rightMargin=2*cm,
                                topMargin=2*cm, bottomMargin=2*cm)
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=16, spaceAfter=20)
        q_style     = ParagraphStyle('Question', parent=styles['Normal'], fontSize=11, fontName='Helvetica-Bold', spaceAfter=6)
        opt_style   = ParagraphStyle('Option', parent=styles['Normal'], fontSize=10, leftIndent=20, spaceAfter=3)
        ans_style   = ParagraphStyle('Answer', parent=styles['Normal'], fontSize=10, leftIndent=20, textColor='green', spaceAfter=12)

        story = [Paragraph("MCQ Question Paper", title_style)]
        for i, m in enumerate(mcqs, 1):
            story.append(Paragraph(f"Q{i}. {m.get('question','')}", q_style))
            for key in ['A','B','C','D']:
                val = m.get(f'option_{key.lower()}','')
                if val:
                    story.append(Paragraph(f"{key}) {val}", opt_style))
            ans = m.get('answer','')
            ans_text = m.get(f"option_{ans.lower()}", ans) if ans else ''
            story.append(Paragraph(f"✓ Answer: {ans}) {ans_text}", ans_style))
            story.append(Spacer(1, 6))

        doc.build(story)
        return {"content": buf.getvalue(), "mime": "application/pdf", "ext": "pdf"}

    if fmt == "docx":
        from docx import Document
        from docx.shared import Pt, RGBColor
        import io

        doc = Document()
        title = doc.add_heading('MCQ Question Paper', 0)
        title.alignment = 1
        doc.add_paragraph('')

        for i, m in enumerate(mcqs, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"Q{i}. {m.get('question', '')}")
            run.bold = True
            run.font.size = Pt(12)
            for label, key in [('A','option_a'),('B','option_b'),('C','option_c'),('D','option_d')]:
                val = m.get(key, '')
                if val:
                    doc.add_paragraph(f"   {label}) {val}")
            ans = m.get('answer', '')
            ans_para = doc.add_paragraph()
            ans_run = ans_para.add_run(f"   ✅ Answer: {ans}")
            ans_run.bold = True
            ans_run.font.color.rgb = RGBColor(0x16, 0xa3, 0x4a)
            doc.add_paragraph('')

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return {
            "content": buf.read(),
            "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "ext": "docx"
        }

    raise ValueError(f"Unsupported export format: {fmt}  (use json, csv, xlsx, pdf or docx)")


def export_from_data(mcqs: list, fmt: str = "json") -> dict:
    if fmt == "json":
        return {"content": json.dumps(mcqs, indent=2, ensure_ascii=False),
                "mime": "application/json", "ext": "json"}

    if fmt == "csv":
        import pandas as pd, io
        buf = io.StringIO()
        pd.DataFrame(mcqs).to_csv(buf, index=False)
        return {"content": buf.getvalue(), "mime": "text/csv", "ext": "csv"}

    if fmt == "xlsx":
        import pandas as pd, io
        buf = io.BytesIO()
        with pd.ExcelWriter(buf, engine="openpyxl") as writer:
            pd.DataFrame(mcqs).to_excel(writer, index=False, sheet_name="MCQs")
        return {"content": buf.getvalue(),
                "mime": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                "ext": "xlsx"}

    if fmt == "pdf":
        import io
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4,
                                leftMargin=2*cm, rightMargin=2*cm,
                                topMargin=2*cm, bottomMargin=2*cm)
        styles = getSampleStyleSheet()
        q_style   = ParagraphStyle('Q',   parent=styles['Normal'], fontSize=11, fontName='Helvetica-Bold', spaceAfter=6)
        opt_style = ParagraphStyle('Opt', parent=styles['Normal'], fontSize=10, leftIndent=20, spaceAfter=3)
        ans_style = ParagraphStyle('Ans', parent=styles['Normal'], fontSize=10, leftIndent=20, textColor='green', spaceAfter=12)

        story = [Paragraph("MCQ Question Paper", styles['Heading1'])]
        for i, m in enumerate(mcqs, 1):
            story.append(Paragraph(f"Q{i}. {m.get('question','')}", q_style))
            for key in ['A','B','C','D']:
                val = m.get(f'option_{key.lower()}','')
                if val:
                    story.append(Paragraph(f"{key}) {val}", opt_style))
            ans = m.get('answer','')
            ans_text = m.get(f"option_{ans.lower()}", ans) if ans else ''
            story.append(Paragraph(f"✓ Answer: {ans}) {ans_text}", ans_style))
            story.append(Spacer(1, 6))
        doc.build(story)
        return {"content": buf.getvalue(), "mime": "application/pdf", "ext": "pdf"}

    if fmt == "docx":
        from docx import Document
        from docx.shared import Pt, RGBColor
        import io

        doc = Document()
        doc.add_heading('MCQ Question Paper', 0).alignment = 1
        doc.add_paragraph('')

        for i, m in enumerate(mcqs, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"Q{i}. {m.get('question','')}")
            run.bold = True
            run.font.size = Pt(12)
            for label, key in [('A','option_a'),('B','option_b'),('C','option_c'),('D','option_d')]:
                val = m.get(key,'')
                if val:
                    doc.add_paragraph(f"   {label}) {val}")
            ans = m.get('answer','')
            ans_para = doc.add_paragraph()
            ans_run = ans_para.add_run(f"   ✅ Answer: {ans}")
            ans_run.bold = True
            ans_run.font.color.rgb = RGBColor(0x16, 0xa3, 0x4a)
            doc.add_paragraph('')

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return {"content": buf.read(),
                "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "ext": "docx"}

    raise ValueError(f"Unsupported format: {fmt}")


# ── SDM Response functions ────────────────────────────────────────────────────

def save_sdm_response(filename: str, item_oid: str, value: str):
    target = filename or _active
    if not target or target not in _index:
        raise ValueError("No active file.")
    if "responses" not in _index[target]:
        _index[target]["responses"] = {}
    _index[target]["responses"][item_oid] = value
    _save(target)
    return {"item_oid": item_oid, "value": value}


def get_sdm_responses(filename: str) -> dict:
    target = filename or _active
    if not target or target not in _index:
        return {}
    return _index[target].get("responses", {})

def delete_sdm_response(filename: str, item_oid: str) -> dict:
    target = filename or _active
    if not target or target not in _index:
        raise ValueError("No active file.")
    responses = _index[target].get("responses", {})
    if item_oid not in responses:
        raise ValueError(f"Response '{item_oid}' not found.")
    del responses[item_oid]
    _index[target]["responses"] = responses
    _save(target)
    return {"deleted": item_oid}

def clear_sdm_responses(filename: str):
    target = filename or _active
    if not target or target not in _index:
        raise ValueError("No active file.")
    _index[target]["responses"] = {}
    _save(target)
    return {"cleared": True}


# ── Compatibility aliases ─────────────────────────────────────────────────────

def add_file(filename: str, content: bytes, original_path: str = "") -> dict:
    return ingest(filename, content, original_path)


def delete_file(filename: str) -> str:
    return remove(filename)


def set_active(filename: str) -> str:
    return switch(filename)


def get_topics(filename: Optional[str] = None) -> list:
    target = filename or _active
    if not target or target not in _index:
        return []
    return sorted(set(
        m.get("topic", "")
        for m in _index[target]["mcqs"]
        if m.get("topic", "")
    ))


def get_stats(filename: Optional[str] = None) -> dict:
    target = filename or _active
    if not target or target not in _index:
        return {}
    mcqs   = _index[target]["mcqs"]
    diffs  = {}
    topics = set()
    for m in mcqs:
        d = m.get("difficulty", "Unknown") or "Unknown"
        diffs[d] = diffs.get(d, 0) + 1
        t = m.get("topic", "")
        if t:
            topics.add(t)
    return {
        "filename":     target,
        "total":        len(mcqs),
        "difficulties": diffs,
        "topics":       sorted(topics),
        "active":       target == _active,
    }